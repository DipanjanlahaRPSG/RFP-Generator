"""
API Routes for RFP Generator
Implements all endpoints for frontend integration
"""

from fastapi import APIRouter, HTTPException, Response
from fastapi.responses import FileResponse
import logging
import json
import uuid
from datetime import datetime
from typing import Dict, Any

from api.schemas import (
    AnalyzeRequest, AnalyzeResponse,
    QuestionRequest, QuestionResponse,
    GenerateRequest, GenerateResponse,
    RegenerateRequest, RegenerateResponse,
    ExportRequest, RFPSections,
    DiscoverContextRequest, DiscoverContextResponse
)
from services.question_generator import QuestionGenerator
from services.section_generator import SectionGenerator
from database.db import get_db

logger = logging.getLogger(__name__)

router = APIRouter()

# Initialize services
question_gen = QuestionGenerator()
section_gen = SectionGenerator()


@router.post("/analyze", response_model=AnalyzeResponse)
async def analyze_rfp_request(request: AnalyzeRequest):
    """
    Analyze initial RFP request and generate follow-up questions
    """
    logger.info(f"Analyzing RFP request: {request.prompt[:100]}...")
    
    try:
        # Create new session
        session_id = str(uuid.uuid4())
        
        # Extract basic entities from prompt
        entities = _extract_entities(request.prompt)
        
        # Determine RFP type
        rfp_type = _classify_rfp_type(request.prompt)
        
        # Generate adaptive questions
        context = {
            "initial_prompt": request.prompt,
            "rfp_type": rfp_type,
            **entities
        }
        questions = question_gen.generate_questions(request.prompt, context)
        
        # Save session to database
        db = await get_db()
        await db.execute(
            """INSERT INTO rfp_sessions (id, title, rfp_type, context)
               VALUES (?, ?, ?, ?)""",
            (session_id, request.prompt[:100], rfp_type, json.dumps(context))
        )
        await db.commit()
        await db.close()
        
        return AnalyzeResponse(
            rfp_type=rfp_type,
            entities=entities,
            questions=questions,
            session_id=session_id
        )
        
    except Exception as e:
        logger.error(f"Error analyzing request: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/questions", response_model=QuestionResponse)
async def get_next_question(request: QuestionRequest):
    """
    Get next question (currently returns completion status)
    """
    # For POC, we generate all questions upfront in /analyze
    # This endpoint just confirms completion
    return QuestionResponse(
        next_question=None,
        is_complete=True
    )


@router.post("/discover-context", response_model=DiscoverContextResponse)
async def discover_rag_context(request: DiscoverContextRequest):
    """
    Search RAG system for relevant historical RFPs before generation
    This step happens after user answers questions
    """
    logger.info(f"Discovering RAG context for session: {request.session_id}")
    
    try:
        from services.rag_context_discovery import RAGContextDiscovery
        
        discovery = RAGContextDiscovery()
        result = await discovery.discover_context(request.context)
        
        # Update session context with discovered insights
        db = await get_db()
        cursor = await db.execute(
            "SELECT context FROM rfp_sessions WHERE id = ?",
            (request.session_id,)
        )
        row = await cursor.fetchone()
        
        if row:
            existing_context = json.loads(row['context'])
            existing_context['rag_discovery'] = result
            
            await db.execute(
                "UPDATE rfp_sessions SET context = ? WHERE id = ?",
                (json.dumps(existing_context), request.session_id)
            )
            await db.commit()
        
        await db.close()
        
        return DiscoverContextResponse(
            session_id=request.session_id,
            **result
        )
        
    except Exception as e:
        logger.error(f"Error discovering context: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/generate", response_model=GenerateResponse)
async def generate_rfp_sections(request: GenerateRequest):
    """
    Generate all 25 RFP sections (NEW + OLD + RULES)
    """
    logger.info(f"Generating RFP sections for session: {request.session_id}")
    
    try:
        # Generate all sections
        sections_dict = section_gen.generate_all_sections(request.context)
        
        # Convert to response format
        sections = RFPSections(
            new=sections_dict["new"],
            old=sections_dict["old"],
            rules=sections_dict["rules"]
        )
        
        # Save sections to database
        db = await get_db()
        for section in sections_dict["new"]:
            await _save_section(db, request.session_id, section, "new")
        for section in sections_dict["old"]:
            await _save_section(db, request.session_id, section, "old")
        for section in sections_dict["rules"]:
            await _save_section(db, request.session_id, section, "rules")
        
        await db.commit()
        await db.close()
        
        return GenerateResponse(
            session_id=request.session_id,
            sections=sections
        )
        
    except Exception as e:
        logger.error(f"Error generating sections: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/regenerate", response_model=RegenerateResponse)
async def regenerate_section(request: RegenerateRequest):
    """
    Regenerate a single section with additional context
    """
    logger.info(f"Regenerating section: {request.section_name} (iteration {request.iteration})")
    
    try:
        # Determine section type
        if request.section_name in section_gen.NEW_SECTIONS:
            section = section_gen.generate_new_section(
                section_name=request.section_name,
                context=request.context,
                iteration=request.iteration,
                additional_context=request.additional_context
            )
        elif request.section_name in section_gen.OLD_SECTIONS:
            section = section_gen.generate_old_section(
                section_name=request.section_name,
                context=request.context,
                iteration=request.iteration
            )
        else:
            # RULES sections can't be regenerated
            raise HTTPException(status_code=400, detail="RULES sections cannot be regenerated")
        
        # Update section in database
        db = await get_db()
        await db.execute(
            """UPDATE sections 
               SET content = ?, ai_eval = ?, regen_count = ?
               WHERE session_id = ? AND name = ?""",
            (
                section["content"],
                json.dumps(section.get("aiEval", {})),
                request.iteration,
                request.session_id,
                request.section_name
            )
        )
        await db.commit()
        await db.close()
        
        return RegenerateResponse(section=section)
        
    except Exception as e:
        logger.error(f"Error regenerating section: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/export")
async def export_rfp(session_id: str):
    """
    Export RFP as Word document
    """
    logger.info(f"Exporting RFP for session: {session_id}")
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Get sections from database
        db = await get_db()
        cursor = await db.execute(
            """SELECT name, content, source_type 
               FROM sections 
               WHERE session_id = ? 
               ORDER BY 
                   CASE source_type 
                       WHEN 'new' THEN 1 
                       WHEN 'old' THEN 2 
                       WHEN 'rules' THEN 3 
                   END,
                   id""",
            (session_id,)
        )
        sections = await cursor.fetchall()
        
        # Get session info
        cursor = await db.execute(
            "SELECT title, rfp_type FROM rfp_sessions WHERE id = ?",
            (session_id,)
        )
        session = await cursor.fetchone()
        await db.close()
        
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading(session['title'] or 'Request for Proposal', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"RFP Type: {session['rfp_type']}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_page_break()
        
        # Add sections
        for section in sections:
            # Add section heading
            doc.add_heading(section['name'], 1)
            
            # Add content (convert markdown to plain text for now)
            content = section['content'].replace('##', '').replace('#', '')
            doc.add_paragraph(content)
            doc.add_paragraph()  # Spacing
        
        # Save to temp file
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        # Return file
        return FileResponse(
            temp_file.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f'RFP_{session_id[:8]}.docx'
        )
        
    except Exception as e:
        logger.error(f"Error exporting RFP: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Helper functions

def _extract_entities(prompt: str) -> Dict[str, Any]:
    """Extract basic entities from prompt"""
    entities = {}
    
    # Extract duration
    import re
    duration_match = re.search(r'(\d+)\s*(month|year|week|day)s?', prompt.lower())
    if duration_match:
        entities['duration'] = f"{duration_match.group(1)} {duration_match.group(2)}s"
    
    # Extract service/role
    # Simple heuristic: look for common keywords
    service_keywords = ['designer', 'developer', 'consultant', 'engineer', 'manager', 'analyst']
    for keyword in service_keywords:
        if keyword in prompt.lower():
            entities['service'] = keyword.title()
            break
    
    return entities


def _classify_rfp_type(prompt: str) -> str:
    """Classify RFP type from prompt"""
    prompt_lower = prompt.lower()
    
    if any(word in prompt_lower for word in ['service', 'consulting', 'design', 'development']):
        return 'Service_Agreement'
    elif any(word in prompt_lower for word in ['supply', 'equipment', 'material', 'procurement']):
        return 'Supply_Contract'
    elif any(word in prompt_lower for word in ['epc', 'construction', 'project', 'infrastructure']):
        return 'EPC_Project'
    else:
        return 'Service_Agreement'  # Default


async def _save_section(db, session_id: str, section: Dict[str, Any], source_type: str):
    """Save section to database"""
    await db.execute(
        """INSERT INTO sections (session_id, name, source_type, content, assumptions, ai_eval)
           VALUES (?, ?, ?, ?, ?, ?)""",
        (
            session_id,
            section["name"],
            source_type,
            section["content"],
            json.dumps(section.get("assumptions", [])),
            json.dumps(section.get("aiEval", {}))
        )
    )
