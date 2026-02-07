from docx import Document
import os
from typing import Dict, Any, Optional
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles extraction and processing of DOCX documents"""

    def __init__(self):
        self.supported_extensions = [".docx", ".doc"]

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract full text content from DOCX file"""
        try:
            doc = Document(file_path)
            full_text = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n".join(full_text)

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise

    def extract_filename_metadata(self, file_path: str) -> Dict[str, str]:
        """Extract metadata from filename patterns"""
        filename = os.path.basename(file_path)
        filename_without_ext = os.path.splitext(filename)[0]

        # Parse filename pattern: DocumentType_Client for Service.docx
        parts = filename_without_ext.split("_")

        metadata = {"filename": filename, "original_filename": filename}

        if len(parts) >= 2:
            # First part is usually document type
            metadata["document_type"] = parts[0]

            # Look for client name (usually contains CESC, CPDL, MPSL)
            for part in parts:
                if any(
                    client in part.upper()
                    for client in ["CESC", "CPDL", "MPSL", "CDPL"]
                ):
                    metadata["client_name"] = part.replace(" ", "_")
                    break

            # Extract service description (after "for")
            if "for" in filename_without_ext.lower():
                service_part = filename_without_ext.lower().split("for")[1]
                metadata["specific_service"] = service_part.replace("_", " ").strip()

        return metadata

    def get_document_stats(self, content: str) -> Dict[str, Any]:
        """Get basic statistics about document content"""
        words = content.split()
        sentences = content.split(".")
        paragraphs = content.split("\n\n")

        return {
            "word_count": len(words),
            "sentence_count": len([s for s in sentences if s.strip()]),
            "paragraph_count": len([p for p in paragraphs if p.strip()]),
            "character_count": len(content),
            "avg_words_per_sentence": len(words) / max(len(sentences), 1),
        }

    def is_supported_file(self, file_path: str) -> bool:
        """Check if file format is supported"""
        _, ext = os.path.splitext(file_path.lower())
        return ext in self.supported_extensions

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Complete document processing pipeline"""
        if not self.is_supported_file(file_path):
            raise ValueError(f"Unsupported file format: {file_path}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Extract text content
        content = self.extract_text_from_docx(file_path)

        # Extract filename metadata
        filename_metadata = self.extract_filename_metadata(file_path)

        # Get document statistics
        stats = self.get_document_stats(content)

        return {
            "content": content,
            "file_path": file_path,
            "filename_metadata": filename_metadata,
            "stats": stats,
            "processing_timestamp": str(datetime.now()),
        }
